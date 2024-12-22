# Required Libraries
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from io import BytesIO
import requests

def add_page_borders(section,check):
    borders = OxmlElement('w:borders')
    if check == True:
        # Define border properties
        for border_type in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_type}')
            border.set(qn('w:val'), 'single')  # Set border style
            border.set(qn('w:sz'), '8')  # Border size (in half-points)
            border.set(qn('w:space'), '0')  # Space around the border
            borders.append(border)
        section.top_margin = Pt(60)
        section.bottom_margin = Pt(60)
        section.left_margin = Pt(60)
        section.right_margin = Pt(60)
        section._sectPr.append(borders)
    else:
        for border_type in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_type}')
            border.set(qn('w:val'), 'single')  # Set border style
            border.set(qn('w:sz'), '10')  # Border size (in half-points)
            border.set(qn('w:space'), '10')  # Space around the border
            borders.append(border)
        section.top_margin = Pt(20)
        section.bottom_margin = Pt(20)
        section.left_margin = Pt(20)
        section.right_margin = Pt(20)
        section._sectPr.append(borders)


    # Add the borders to the section's properties


def set_vertical_alignment(cell, alignment):
    """
    Set the vertical alignment for a table cell.

    Parameters:
    - cell: The cell object from python-docx.
    - alignment: The desired alignment ('top', 'center', 'bottom').
    """
    # Get the cell's properties
    cell_properties = cell._element.get_or_add_tcPr()

    # Create a new vertical alignment element
    v_align = OxmlElement('w:vAlign')
    v_align.set(qn('w:val'), alignment)

    # Append the vertical alignment to the cell properties
    cell_properties.append(v_align)

def set_column_widths(table, column_indices, width_in_inches):
    """
    Set custom widths for specified columns in a table.

    Parameters:
    - table: The table object from python-docx.
    - column_indices: List of column indices to adjust.
    - width_in_inches: The desired width in inches for the specified columns.
    """
    for row in table.rows:
        for col_idx in column_indices:
            if col_idx < len(row.cells):
                cell = row.cells[col_idx]
                cell_width = cell._element
                tc_width = cell_width.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(int(width_in_inches * 1440)))  # Inches to twips
                tcW.set(qn('w:type'), 'dxa')
                tc_width.append(tcW)

def generate_report(file_path, sheet_name, column_start, column_end):
    """
    Generate a student analysis report in .docx format from an Excel sheet.

    Parameters:
    - file_path: Path to the Excel file.
    - sheet_name: Name of the sheet corresponding to the student ID.
    - column_start: Starting column (e.g., 'A').
    - column_end: Ending column (e.g., 'H').

    Returns:
    - Path to the generated .docx report.
    """
    # Load Excel file
    try:
        xls = pd.ExcelFile(file_path)
    except Exception as e:
        print(f"Error loading Excel file: {e}")
        return None

    # Validate sheet name
    if sheet_name not in xls.sheet_names or sheet_name.upper() == "RESOURCE":
        print("Invalid Sheet Name or it's a RESOURCE sheet.")
        return None

    # Load the specified sheet without headers
    try:
        df = pd.read_excel(file_path, sheet_name=sheet_name, usecols=f"{column_start}:{column_end}", header=None)
    except Exception as e:
        print(f"Error reading the sheet: {e}")
        return None

    # Initialize Document
    doc = Document()

    # Variables to track the current position
    row_idx = 0
    num_rows = df.shape[0]

    # FRONT PAGE: Student Details
    student_details = {}
    if row_idx + 1 < num_rows:
        headers = df.iloc[row_idx].tolist()
        values = df.iloc[row_idx + 1].tolist()

        for header, value in zip(headers, values):
            if header == "EXAM TYPE":
                fileName = value
            if pd.notna(header) and pd.notna(value):
                student_details[header.strip().upper()] = value
        row_idx += 3  # Move past headers and values


    for _ in range(5):  # Adjust the number of paragraphs as needed
        doc.add_paragraph()  # Add empty paragraph

    # Add Front Page Title
    title = doc.add_paragraph()
    run = title.add_run("STUDENT ANALYSIS REPORT")
    run.font.size = Pt(24)
    run.bold = True
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Insert Logo if available
    logo_url = student_details.get("LOGO")
    if logo_url:
        try:
            # Modify the URL to access the raw image content if it's a Google Drive link
            if "drive.google.com" in logo_url:
                file_id = logo_url.split('/d/')[1].split('/')[0]
                logo_url = f"https://drive.google.com/uc?id={file_id}"

            response = requests.get(logo_url)
            if response.status_code == 200:
                image = BytesIO(response.content)
                doc.add_picture(image, width=Inches(2))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        except Exception as e:
            print("Error loading logo:", e)

    # Add Student Details
    details_paragraph = doc.add_paragraph()
    details_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    details_run = details_paragraph.add_run()

    student_info_keys = ["ID", "NAME", "STANDARD", "SECTION", "EXAM TYPE", "SCHOOL NAME"]
    details_text = ""
    for key in student_info_keys:
        value = student_details.get(key.upper(), "N/A")
        details_text += f"{key}: {value}\n"

    details_run.text = details_text.strip()
    details_run.font.size = Pt(14)
    details_run.font.bold = False

    # Add some spacing
    doc.add_paragraph().add_run("\n")

    # SECOND PAGE: Overall Summary Table
    # Add page break
    doc.add_page_break()

    for _ in range(6):  # Adjust the number of paragraphs as needed
        doc.add_paragraph()  # Add empty paragraph


    # Overall Summary Heading
    overall_heading = doc.add_heading("OVERALL PERFORMANCE SUMMARY", level=1)
    run = overall_heading.runs[0]
    run.font.color.rgb = RGBColor(31, 73, 125)  # Set the color (e.g., blue)
    run.font.name = 'Calibri'  # Set font style (e.g., Arial)
    run.font.size = Pt(18)  # Set font size
    run.font.bold = True  # Set bold if desired

    # Ensure font style is applied correctly
    overall_heading.style.font.name = 'Calibri'
    overall_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Read Overall Summary Table Headers
    if row_idx < num_rows:
        overall_headers = df.iloc[row_idx].dropna().tolist()
        row_idx += 1
    else:
        overall_headers = []

    for _ in range(2):  # Adjust the number of paragraphs as needed
        doc.add_paragraph()  # Add empty paragraph

    # Create Overall Summary Table
    if overall_headers:
        overall_table = doc.add_table(rows=1, cols=len(overall_headers))
        overall_table.style = 'Colorful List Accent 2'

        overall_table.autofit = False

        # Set the width of the table (in inches)
        table_width = Inches(6)  # Adjust as needed
        overall_table.width = table_width

        # Set widths for each column (adjust as needed)
        for column in overall_table.columns:
            column.width = Inches(2)

        hdr_cells = overall_table.rows[0].cells
        for i, header in enumerate(overall_headers):
            hdr_cells[i].text = str(header).strip()
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            #tttt
        overall_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for row in overall_table.rows:
            for cell in row.cells:
                set_vertical_alignment(cell, 'center')#added
                # Clear existing paragraphs and create a new one if necessary
                if not cell.paragraphs:
                    cell.add_paragraph("")  # Ensure there's at least one paragraph

                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(14)

                    # Optionally, if you want to set the font size for text added later
                    # This ensures that any new text added to the paragraph will also have the font size
                    if paragraph.text:  # Only apply if there's text
                        run = paragraph.add_run()  # Add a new run if needed
                        run.font.size = Pt(14)


        # Fill Overall Summary Table with Data
        while row_idx < num_rows and not df.iloc[row_idx].isna().all():
            row = df.iloc[row_idx].tolist()
            # Ensure the row has the same number of columns as headers
            if len(row) < len(overall_headers):
                row += [""] * (len(overall_headers) - len(row))
            row = row[:len(overall_headers)]
            # Add row to table
            row_cells = overall_table.add_row().cells

            for i, cell_value in enumerate(row):
                row_cells[i].text = str(cell_value).strip() if pd.notna(cell_value) else ""
                # Align numeric columns to center or right if needed

                row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                set_vertical_alignment(row_cells[i], 'center') #added
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(14)
            row_idx += 1

    # Skip any empty rows
    while row_idx < num_rows and df.iloc[row_idx].isna().all():
        row_idx += 1

    # SUBJECT PAGES: Unit-wise and Topic-wise tables
    while row_idx < num_rows:
        current_row = df.iloc[row_idx].tolist()

        # Check if the row indicates a new subject
        if pd.notna(current_row[0]) and str(current_row[0]).strip() == "Subject":
            # Get Subject Name
            subject_name = current_row[1] if len(current_row) > 1 and pd.notna(current_row[1]) else "Unknown Subject"

            # Add page break for each subject
            doc.add_page_break()

            for _ in range(0):  # Adjust the number of paragraphs as needed
                doc.add_paragraph()  # Add empty paragraph

            # Subject Heading
            subject_heading = doc.add_heading(f"{subject_name}", level=1)
            run = subject_heading.runs[0]
            run.font.color.rgb = RGBColor(31, 73, 125)  # Set the color (e.g., blue)
            run.font.name = 'Calibri'  # Set font style (e.g., Arial)
            run.font.size = Pt(18)  # Set font size
            run.font.bold = True  # Set bold if desired

            # Ensure font style is applied correctly
            subject_heading.style.font.name = 'Calibri'
            subject_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            for _ in range(1):  # Adjust the number of paragraphs as needed
                doc.add_paragraph()  # Add empty paragraph

            row_idx += 1  # Move to Unit-wise Report heading
            # UNIT-WISE REPORT light
            if row_idx < num_rows:
                # Add Unit-wise Heading
                unit_heading = doc.add_heading("UNIT-WISE ANALYSIS", level=2)
                run = unit_heading.runs[0]
                run.font.color.rgb = RGBColor(31, 73, 125)  # Set the color (e.g., blue)
                run.font.name = 'Calibri'  # Set font style (e.g., Arial)
                run.font.size = Pt(14)  # Set font size
                run.font.bold = True  # Set bold if desired

                # Ensure font style is applied correctly
                unit_heading.style.font.name = 'Calibri'
                unit_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                for _ in range(0):  # Adjust the number of paragraphs as needed
                    doc.add_paragraph()  # Add empty paragraph

                # Read Unit-wise Table Headers
                unit_headers = df.iloc[row_idx].dropna().tolist()
                row_idx += 1

                if unit_headers:
                    unit_table = doc.add_table(rows=1, cols=len(unit_headers))
                    unit_table.style = 'Medium Shading 1 Accent 3'
                    hdr_cells = unit_table.rows[0].cells
                    for i, header in enumerate(unit_headers):
                        hdr_cells[i].text = str(header).strip()
                        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                        hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    #ttt
                    unit_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    # Reduce font size for table content
                    for row in unit_table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(9)

                    # Identify the index of "OBTAINED MARKS" and "PERFORMANCE %" columns
                    try:
                        obtained_marks_idx = unit_headers.index("OBTAINED MARKS")
                        performance_idx = unit_headers.index("PERFORMANCE %")
                        overall_idx = unit_headers.index("OVERALL %")
                    except ValueError:
                        obtained_marks_idx = -1  # Column not found
                        performance_idx = -1  # Column not found
                        overall_idx = -1

                    # Fill Unit-wise Table with Data
                    while row_idx < num_rows and not df.iloc[row_idx].isna().all():
                        unit_row = df.iloc[row_idx].tolist()
                        if len(unit_row) < len(unit_headers):
                            unit_row += [""] * (len(unit_headers) - len(unit_row))
                        unit_row = unit_row[:len(unit_headers)]

                        # Check "OBTAINED MARKS" value
                        if obtained_marks_idx != -1:
                            obtained_marks = str(unit_row[obtained_marks_idx]).strip()
                            if obtained_marks == "0/0":
                                row_idx += 1
                                continue  # Skip this row


                        # Add row to table
                        row_cells = unit_table.add_row().cells
                        for i, cell_value in enumerate(unit_row):
                            row_cells[i].text = str(cell_value).strip() if pd.notna(cell_value) else ""
                            # Align numeric columns to center or right if needed
                            row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            # Change background color for "PERFORMANCE %" column if below 50
                            if overall_idx != -1 and i == overall_idx:
                                if isinstance(cell_value, (int, float)) and cell_value < 50:
                                    cell = row_cells[i]._element
                                    cell_properties = cell.get_or_add_tcPr()
                                    cell_shading = OxmlElement('w:shd')
                                    cell_shading.set(qn('w:fill'), '#FF8585')  # Red color
                                    cell_properties.append(cell_shading)

                        row_idx += 1

                    # Adjust column widths for specific columns
                    if obtained_marks_idx != -1:
                        set_column_widths(unit_table, [obtained_marks_idx], 1.0)  # Adjust "OBTAINED MARKS"

                # Skip the single empty row between Unit-wise and Topic-wise tables
                while row_idx < num_rows and df.iloc[row_idx].isna().all():
                    row_idx += 1

                # TOPIC-WISE REPORT
                if row_idx < num_rows:
                    # Check if the Topic-wise section starts at the end of the page
                    if len(doc.paragraphs) > 0 and doc.paragraphs[-1].text.strip() == "":
                        doc.add_page_break()  # Add a page break if the last paragraph is not empty\
                    for _ in range(1):  # Adjust the number of paragraphs as needed
                        doc.add_paragraph()  # Add empty paragraph
                    # Add Topic-wise Heading
                    topic_heading = doc.add_heading("TOPIC-WISE ANALYSIS", level=2)
                    run = topic_heading.runs[0]
                    run.font.color.rgb = RGBColor(31, 73, 125)  # Set the color (e.g., blue)
                    run.font.name = 'Calibri'  # Set font style (e.g., Arial)
                    run.font.size = Pt(14)  # Set font size
                    run.font.bold = True  # Set bold if desired

                    # Ensure font style is applied correctly
                    topic_heading.style.font.name = 'Calibri'
                    topic_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    for _ in range(0):  # Adjust the number of paragraphs as needed
                        doc.add_paragraph()  # Add empty paragraph

                    # Read Topic-wise Table Headers
                    topic_headers = df.iloc[row_idx].dropna().tolist()
                    row_idx += 1

                    if topic_headers:
                        topic_table = doc.add_table(rows=1, cols=len(topic_headers))
                        topic_table.style = 'Medium Shading 1 Accent 4'
                        hdr_cells = topic_table.rows[0].cells
                        for i, header in enumerate(topic_headers):
                            hdr_cells[i].text = str(header).strip()
                            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                            hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        #ttt
                        topic_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        # Reduce font size for table content
                        for row in topic_table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.size = Pt(9)

                        # Identify the index of "OBTAINED MARKS" and "PERFORMANCE %" columns
                        try:
                            obtained_marks_idx_topic = topic_headers.index("OBTAINED MARKS")
                            performance_idx_topic = topic_headers.index("PERFORMANCE %")
                            overall_idx_topic = topic_headers.index("OVERALL %")
                        except ValueError:
                            obtained_marks_idx_topic = -1  # Column not found
                            performance_idx_topic = -1  # Column not found
                            overall_idx_topic = -1


                        # Fill Topic-wise Table with Data
                        while row_idx < num_rows and not df.iloc[row_idx].isna().all():
                            topic_row = df.iloc[row_idx].tolist()
                            if len(topic_row) < len(topic_headers):
                                topic_row += [""] * (len(topic_headers) - len(topic_row))
                            topic_row = topic_row[:len(topic_headers)]

                            # Check "OBTAINED MARKS" value
                            if obtained_marks_idx_topic != -1:
                                obtained_marks_topic = str(topic_row[obtained_marks_idx_topic]).strip()
                                if obtained_marks_topic == "0/0":
                                    row_idx += 1
                                    continue  # Skip this row

                            # Add row to table
                            row_cells = topic_table.add_row().cells
                            for i, cell_value in enumerate(topic_row):
                                row_cells[i].text = str(cell_value).strip() if pd.notna(cell_value) else ""
                                row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                                # Change background color for "PERFORMANCE %" column if below 50
                                if overall_idx_topic != -1 and i == overall_idx_topic:
                                    if isinstance(cell_value, (int, float)) and cell_value < 50:
                                        cell = row_cells[i]._element
                                        cell_properties = cell.get_or_add_tcPr()
                                        cell_shading = OxmlElement('w:shd')
                                        cell_shading.set(qn('w:fill'), '#FF8585')  # Red color
                                        cell_properties.append(cell_shading)

                            row_idx += 1

                        # Adjust column widths for specific columns
                        if obtained_marks_idx_topic != -1:
                            set_column_widths(topic_table, [obtained_marks_idx_topic], 1.0)  # Adjust "OBTAINED MARKS"

        elif pd.notna(current_row[0]) and str(current_row[0]).strip().upper() == "END":
            # Terminate processing if "END" is found
            break
        else:
            # Move to next row if current row doesn't indicate a subject
            row_idx += 1

        # Skip any remaining empty rows before the next subject
        while row_idx < num_rows and df.iloc[row_idx].isna().all():
            row_idx += 1
    add_page_borders(doc.sections[-1], False)
    # Save Document
    output_path = f"/{fileName}_Report.docx"
    try:
        doc.save(output_path)
        print(f"Report successfully generated: {output_path}")
        return output_path
    except Exception as e:
        print(f"Error saving the document: {e}")
        return None

# Execution Block
if __name__ == "__main__":
    import os
    # Prompt user for the file path
    file_path = r"C:\Users\Sanjay\student_data.xlsx"  # Use raw string (r"") for Windows paths

    if not os.path.exists(file_path):
        print("The specified file path does not exist. Please check and try again.")
    else:
        # Prompt user for the Student ID (which corresponds to the sheet name)
        student_id = "1001"  # Replace with user input or predefined value

        # Define the column range (e.g., A-H). Adjust as needed for different exams.
        column_start = "AC"
        column_end = "AJ"

        # Generate the report
        report_path = generate_report(file_path, sheet_name=student_id, column_start=column_start, column_end=column_end)

        # If report is generated successfully, inform the user
        if report_path and os.path.exists(report_path):
            print(f"Report generated successfully. You can find it here: {report_path}")
        else:
            print("Failed to generate the report.")